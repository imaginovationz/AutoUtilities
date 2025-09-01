'''
Created on Aug 3, 2025

@author: nigam
'''
import docx
import json
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_alignment_name(alignment):
    if alignment == WD_ALIGN_PARAGRAPH.LEFT:
        return 'left'
    elif alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return 'center'
    elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return 'right'
    elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return 'justify'
    elif alignment == WD_ALIGN_PARAGRAPH.DISTRIBUTE:
        return 'distribute'
    else:
        return 'unknown'

def get_effective_font_name(run):
    if run.font and run.font.name:
        return run.font.name
    style = getattr(run, 'style', None)
    while style:
        font_name = getattr(getattr(style, 'font', None), 'name', None)
        if font_name:
            return font_name
        style = getattr(style, 'base_style', None)
    para_style = getattr(run, '_paragraph', None)
    style = getattr(para_style, 'style', None)
    while style:
        font_name = getattr(getattr(style, 'font', None), 'name', None)
        if font_name:
            return font_name
        style = getattr(style, 'base_style', None)
    return None

def get_document_default_font(doc):
    try:
        normal_style = doc.styles['Normal']
        return normal_style.font.name
    except Exception:
        return None

def extract_table_content(table, doc, table_index, tables_content, parent_table_indices=None):
    if parent_table_indices is None:
        parent_table_indices = []
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_paragraphs = []
            for para in cell.paragraphs:
                alignment = get_alignment_name(para.alignment)
                para_style = para.style.name if para.style else None
                runs_info = []
                for run in para.runs:
                    font_name = get_effective_font_name(run)
                    if not font_name:
                        font_name = get_document_default_font(doc)
                    run_info = {
                        'text': run.text,
                        'font_name': font_name,
                        'font_size': run.font.size.pt if run.font.size else None,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
                        'style': run.style.name if run.style else None
                    }
                    runs_info.append(run_info)
                cell_paragraphs.append({
                    'text': para.text.strip(),
                    'alignment': alignment,
                    'style': para_style,
                    'runs': runs_info
                })
            # Extract nested tables in the cell
            for nested_table in cell.tables:
                nested_table_indices = parent_table_indices + [table_index]
                extract_table_content(nested_table, doc, 0, tables_content, parent_table_indices=nested_table_indices)
            row_data.append(cell_paragraphs)
        table_data.append(row_data)
    tables_content.append({
        'type': 'table',
        'table_index': table_index if not parent_table_indices else parent_table_indices + [table_index],
        'data': table_data
    })

def extract_all_tables_with_formatting(file_path):
    doc = docx.Document(file_path)
    tables_content = []
    for table_index, table in enumerate(doc.tables):
        extract_table_content(table, doc, table_index, tables_content)
    return tables_content

# Sample usage:
if __name__ == "__main__":
    file = '..//DocParser_AI//Files//PS00082.docx'
    tables = extract_all_tables_with_formatting(file)
    # Output as JSON
    print(json.dumps(tables, ensure_ascii=False, indent=2))