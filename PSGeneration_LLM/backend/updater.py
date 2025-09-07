import os
import logging
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from utils import UPDATED_DIR
from generate_txt_from_docx import save_chunks_info

logger = logging.getLogger(__name__)

def normalize(text):
    return ' '.join(text.strip().split()).lower()

def build_ps_structure_map(doc):
    structure = {}
    current_section = []
    para_counter = {}
    last_heading = None
    for idx, para in enumerate(doc.paragraphs):
        style = para.style.name if hasattr(para.style, 'name') else str(para.style)
        if style.replace(" ", "").lower().startswith("heading"):
            try:
                level = int(''.join(filter(str.isdigit, style)))
            except Exception:
                level = 1
            heading_text = para.text.strip()
            last_heading = heading_text if heading_text else last_heading
            current_section = current_section[:level-1] + [heading_text]
        elif not current_section and para.text.strip():
            last_heading = "ROOT"
            current_section = [last_heading]
        section_path = " > ".join([s for s in current_section if s])
        if section_path not in para_counter:
            para_counter[section_path] = 0
        para_idx = para_counter[section_path]
        para_counter[section_path] += 1
        if section_path not in structure:
            structure[section_path] = []
        structure[section_path].append({
            "para_idx": para_idx,
            "docx_idx": idx,
            "style": style,
            "text": para.text
        })
    return structure

def find_anchor_paragraph(structure_map, section_path, para_idx):
    # fallback to "ROOT" if section_path is empty or not found
    if not section_path or section_path not in structure_map:
        if "ROOT" in structure_map:
            section_path = "ROOT"
        else:
            # fallback to first available section
            section_path = next(iter(structure_map.keys()))
    for entry in structure_map[section_path]:
        if entry["para_idx"] == para_idx:
            return entry["docx_idx"], entry["style"]
    # fallback: last para of section
    return structure_map[section_path][-1]["docx_idx"], structure_map[section_path][-1]["style"]

def copy_format(source_para, target_para):
    """Copies formatting from source_para to target_para (font name, size, bold, etc)."""
    try:
        if source_para.runs and target_para.runs:
            target_run = target_para.runs[0]
            source_run = source_para.runs[0]
            target_run.font.name = source_run.font.name
            target_run.font.bold = source_run.font.bold
            target_run.font.size = source_run.font.size
            target_run.font.italic = source_run.font.italic
            target_run.font.underline = source_run.font.underline
    except Exception as e:
        logger.warning(f"Format copy failed: {e}")

def update_ps_document_closest(
    original_path,
    new_to_old_mockup_matches,  # [{requirement_chunk, matched_old_mockup_chunk, similarity, ...}]
    old_mockup_to_ps_matches,   # [{old_mockup_chunk, matched_ps_chunk, similarity, ps_metadata}]
    similarity_threshold=0.7,
    job_id=None, set_progress=None
):
    """
    Deterministic: For each new mockup chunk,
    - Find best-matching Old Mockup chunk (by embedding)
    - Use that to find best anchor section in Old PS (by embedding)
    - If similarity >= threshold:
        - If anchor para is semantically similar: update/replace
        - Else: insert as new para after anchor
    - Copy style from anchor para.
    - Never use LLM to generate/merge.
    """
    status_updates = []
    status_updates.append("Loading original PS document for update.")

    try:
        doc = Document(original_path)
    except Exception as e:
        logger.error(f"Failed to load document: {e}")
        status_updates.append(f"Error loading document: {e}")
        if set_progress and job_id:
            set_progress(job_id, 100, "Error loading document")
        return None, status_updates

    tmp_path = os.path.join(UPDATED_DIR, "tmp_clone.docx")
    with open(original_path, "rb") as src, open(tmp_path, "wb") as dst:
        dst.write(src.read())
    new_doc = Document(tmp_path)
    status_updates.append("Cloned PS document for safe editing.")

    # --- Build section/para structure map for PS ---
    structure_map = build_ps_structure_map(new_doc)

    updated_count = 0
    inserted_count = 0
    total = len(new_to_old_mockup_matches)

    # For detailed logging
    chunk_logs = []

    for idx, m in enumerate(new_to_old_mockup_matches):
        new_chunk = m.get("requirement_chunk", "").strip()
        matched_old_mockup_chunk = m.get("matched_old_mockup_chunk", "").strip()
        mockup_similarity = m.get("similarity", 0.0)
        old_mockup_idx = None
        # Find the best anchor in PS (via the matched old mockup chunk)
        for i, om in enumerate(old_mockup_to_ps_matches):
            if normalize(om.get("old_mockup_chunk", "")) == normalize(matched_old_mockup_chunk):
                old_mockup_idx = i
                break

        action_taken = None
        anchor_section_path = anchor_para_idx = anchor_docx_idx = None
        anchor_ps_chunk = ""
        anchor_ps_similarity = 0.0

        if old_mockup_idx is None:
            status_updates.append(f"Could not find matching Old Mockup for New Mockup chunk {idx+1}. Skipping.")
            action_taken = "No anchor found - skipped"
        else:
            anchor_match = old_mockup_to_ps_matches[old_mockup_idx]
            anchor_ps_similarity = anchor_match.get("similarity", 0.0)
            anchor_ps_similarity = min(max(anchor_ps_similarity, 0), 1)
            ps_meta = anchor_match.get("ps_metadata", {}) or {}
            anchor_section_path = ps_meta.get("section_path", "")
            anchor_para_idx = ps_meta.get("para_idx", 0)
            anchor_ps_chunk = anchor_match.get("matched_ps_chunk", "")

            if anchor_section_path == "" or anchor_para_idx is None:
                logger.warning(f"Anchor metadata missing for chunk {idx+1}: ps_meta={ps_meta}")

            if anchor_ps_similarity is None or anchor_ps_similarity < similarity_threshold:
                status_updates.append(f"Best PS anchor similarity for chunk {idx+1} is {anchor_ps_similarity if anchor_ps_similarity is not None else 0.0:.2f} (< {similarity_threshold}), skipping.")
                action_taken = "Similarity below threshold - skipped"
            else:
                anchor_docx_idx, anchor_style = find_anchor_paragraph(structure_map, anchor_section_path, anchor_para_idx)
                if anchor_docx_idx is None:
                    status_updates.append(f"Could not find anchor location in PS for chunk {idx+1}, skipping.")
                    action_taken = "Anchor location not found - skipped"
                else:
                    para = new_doc.paragraphs[anchor_docx_idx]
                    if normalize(para.text) == normalize(anchor_ps_chunk):
                        para.clear()
                        run = para.add_run(new_chunk)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        copy_format(para, para)
                        updated_count += 1
                        status_updates.append(f"Replaced anchor para {anchor_docx_idx} for chunk {idx+1}.")
                        action_taken = f"Replaced at section_path='{anchor_section_path}' para_idx={anchor_para_idx} (docx_idx={anchor_docx_idx})"
                    else:
                        # Insert after anchor
                        new_para = new_doc.add_paragraph("")
                        p = new_doc.paragraphs.pop()
                        new_doc.paragraphs.insert(anchor_docx_idx + 1, p)
                        p.clear()
                        run = p.add_run(new_chunk)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        copy_format(para, p)
                        inserted_count += 1
                        status_updates.append(f"Inserted new para after anchor {anchor_docx_idx} for chunk {idx+1}.")
                        action_taken = f"Inserted after section_path='{anchor_section_path}' para_idx={anchor_para_idx} (docx_idx={anchor_docx_idx})"

        if set_progress and job_id:
            percent = int(10 + 80 * (idx + 1) / max(1, total))
            set_progress(job_id, percent, f"Inserting chunk {idx+1} of {total}...")

        chunk_logs.append({
            "requirement_chunk": new_chunk,
            "matched_old_mockup_chunk": matched_old_mockup_chunk,
            "matched_ps_chunk": anchor_ps_chunk,
            "mockup_similarity": mockup_similarity,
            "ps_similarity": anchor_ps_similarity,
            "anchor_section_path": anchor_section_path,
            "anchor_para_idx": anchor_para_idx,
            "anchor_docx_idx": anchor_docx_idx,
            "action_taken": action_taken
        })

    out_name = os.path.splitext(os.path.basename(original_path))[0] + "_new_generated.docx"
    out_path = os.path.join(UPDATED_DIR, out_name)
    try:
        new_doc.save(out_path)
        status_updates.append(
            f"Saving new PS document: {out_name} ({updated_count} sections updated, {inserted_count} inserted)."
        )
        status_updates.append("Generation complete.")
        if set_progress and job_id:
            set_progress(job_id, 100, "Generation complete.")
    except Exception as e:
        logger.error(f"Error saving new PS document: {e}")
        status_updates.append(f"Error saving new PS document: {e}")
        if set_progress and job_id:
            set_progress(job_id, 100, "Error saving file")
        return None, status_updates

    doc_id = "ps_new_chunks"
    out_txt_path = save_chunks_info(doc_id, chunk_logs, status_updates)

    return out_path, status_updates