'''
Created on Jul 29, 2025

@author: nigam
'''
import docx
import os
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json #output in json

def get_alignment_name(alignment):
    if alignment == WD_ALIGN_PARAGRAPH.LEFT:
        return 'left'
    elif alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return 'center'
    elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return 'right'
    elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return 'justify'
    elif alignment == WD_ALIGN_PARAGRAPH.DISTRIBUTE:
        return 'distribute'
    else:
        return 'unknown'

def get_effective_font_size(run, doc):
    if run.font and run.font.size:
        return run.font.size.pt
    style = getattr(run, 'style', None)
    while style:
        font_size = getattr(getattr(style, 'font', None), 'size', None)
        if font_size:
            return font_size.pt
        style = getattr(style, 'base_style', None)
    para_style = getattr(run, '_paragraph', None)
    style = getattr(para_style, 'style', None)
    while style:
        font_size = getattr(getattr(style, 'font', None), 'size', None)
        if font_size:
            return font_size.pt
        style = getattr(style, 'base_style', None)
    try:
        normal_style = doc.styles['Normal']
        if normal_style.font.size:
            return normal_style.font.size.pt
    except Exception:
        pass
    return None

def get_effective_font_name(run):
    if run.font and run.font.name:
        return run.font.name
    style = getattr(run, 'style', None)
    while style:
        font_name = getattr(getattr(style, 'font', None), 'name', None)
        if font_name:
            return font_name
        style = getattr(style, 'base_style', None)
    para_style = getattr(run, '_paragraph', None)
    style = getattr(para_style, 'style', None)
    while style:
        font_name = getattr(getattr(style, 'font', None), 'name', None)
        if font_name:
            return font_name
        style = getattr(style, 'base_style', None)
    return None

def get_document_default_font(doc):
    try:
        normal_style = doc.styles['Normal']
        return normal_style.font.name
    except Exception:
        return None

def get_effective_italic(run, doc):
    if run.italic is not None:
        return run.italic
    style = getattr(run, 'style', None)
    while style:
        italic = getattr(getattr(style, 'font', None), 'italic', None)
        if italic is not None:
            return italic
        style = getattr(style, 'base_style', None)
    para_style = getattr(run, '_paragraph', None)
    style = getattr(para_style, 'style', None)
    while style:
        italic = getattr(getattr(style, 'font', None), 'italic', None)
        if italic is not None:
            return italic
        style = getattr(style, 'base_style', None)
    try:
        normal_style = doc.styles['Normal']
        if normal_style.font.italic is not None:
            return normal_style.font.italic
    except Exception:
        pass
    return None

def get_effective_bold(run, doc):
    if run.bold is not None:
        return run.bold
    style = getattr(run, 'style', None)
    while style:
        bold = getattr(getattr(style, 'font', None), 'bold', None)
        if bold is not None:
            return bold
        style = getattr(style, 'base_style', None)
    para_style = getattr(run, '_paragraph', None)
    style = getattr(para_style, 'style', None)
    while style:
        bold = getattr(getattr(style, 'font', None), 'bold', None)
        if bold is not None:
            return bold
        style = getattr(style, 'base_style', None)
    try:
        normal_style = doc.styles['Normal']
        if normal_style.font.bold is not None:
            return normal_style.font.bold
    except Exception:
        pass
    return None

def get_effective_underline(run, doc):
    if run.underline is not None:
        return run.underline
    style = getattr(run, 'style', None)
    while style:
        underline = getattr(getattr(style, 'font', None), 'underline', None)
        if underline is not None:
            return underline
        style = getattr(style, 'base_style', None)
    para_style = getattr(run, '_paragraph', None)
    style = getattr(para_style, 'style', None)
    while style:
        underline = getattr(getattr(style, 'font', None), 'underline', None)
        if underline is not None:
            return underline
        style = getattr(style, 'base_style', None)
    try:
        normal_style = doc.styles['Normal']
        if normal_style.font.underline is not None:
            return normal_style.font.underline
    except Exception:
        pass
    return None

def get_effective_color(run, doc):
    if run.font and run.font.color and run.font.color.rgb:
        return str(run.font.color.rgb)
    style = getattr(run, 'style', None)
    while style:
        color = getattr(getattr(style, 'font', None), 'color', None)
        if color and color.rgb:
            return str(color.rgb)
        style = getattr(style, 'base_style', None)
    para_style = getattr(run, '_paragraph', None)
    style = getattr(para_style, 'style', None)
    while style:
        color = getattr(getattr(style, 'font', None), 'color', None)
        if color and color.rgb:
            return str(color.rgb)
        style = getattr(style, 'base_style', None)
    try:
        normal_style = doc.styles['Normal']
        if normal_style.font.color and normal_style.font.color.rgb:
            return str(normal_style.font.color.rgb)
    except Exception:
        pass
    return None

def parse_docx(file_path):
    doc = docx.Document(file_path)
    parsed_content = []

    for para in doc.paragraphs:
        content_type = 'heading' if para.style.name.startswith('Heading') else 'paragraph'
        alignment = get_alignment_name(para.alignment)
        runs_info = []
        for run in para.runs:
            font_name = get_effective_font_name(run)
            if not font_name:
                font_name = get_document_default_font(doc)
            run_info = {
                'text': run.text,
                'font_name': font_name,
                'font_size': get_effective_font_size(run, doc),
                'bold': get_effective_bold(run, doc),
                'italic': get_effective_italic(run, doc),
                'underline': get_effective_underline(run, doc),
                'color': get_effective_color(run, doc)
            }
            runs_info.append(run_info)
        parsed_content.append({
            'type': content_type,
            'text': para.text.strip(),
            'style': para.style.name,
            'alignment': alignment,
            'runs': runs_info
        })

    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        parsed_content.append({
            'font_name': font_name,
            'type': 'table',
            'table_index': table_idx,
            'data': table_data
        })

    embedded_files = []
    for rel in doc.part.rels.values():
        if rel.reltype == RT.OLE_OBJECT:
            embedded_files.append(rel.target_ref)
        elif rel.reltype == RT.PACKAGE:
            embedded_files.append(rel.target_ref)
    if embedded_files:
        parsed_content.append({
            'type': 'embedded_files',
            'files': embedded_files
        })
    return parsed_content

# Sample usage
if __name__ == "__main__":
    #file = '..//DocParser_AI//Files//PS00082.docx'
    if not os.path.exists(file):
        print(f"Sample file '{file}' not found.")
    else:
        parsed = parse_docx(file)
        print("Parsed Content:")
        for item in parsed:
            print(json.dumps(parsed, ensure_ascii=False, indent=2))