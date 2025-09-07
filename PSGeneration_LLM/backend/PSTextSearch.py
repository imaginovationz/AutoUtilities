from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def highlight_in_paragraph(para, search_text, matches, location_desc):
    """
    Search for search_text inside a paragraph and highlight it,
    even if the text spans multiple runs.
    """
    full_text = "".join(run.text for run in para.runs)
    if search_text not in full_text:
        return

    matches.append({
        "location": location_desc,
        "text": full_text.strip()
    })

    # Reconstruct across runs
    idx = full_text.find(search_text)
    start, end = idx, idx + len(search_text)

    pos = 0
    for run in para.runs:
        run_len = len(run.text)
        run_start, run_end = pos, pos + run_len

        # If this run overlaps with the match, highlight it
        if run_start < end and run_end > start:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        pos += run_len


def search_and_highlight(docx_path, search_text, output_path):
    doc = Document(docx_path)
    matches = []

    def search_in_table(table, path_prefix):
        for r_index, row in enumerate(table.rows, start=1):
            for c_index, cell in enumerate(row.cells, start=1):
                cell_path = f"{path_prefix} -> Row {r_index}, Col {c_index}"

                for p_index, para in enumerate(cell.paragraphs, start=1):
                    highlight_in_paragraph(
                        para, search_text, matches,
                        f"{cell_path} -> Paragraph {p_index}"
                    )

                for t_index, nested_table in enumerate(cell.tables, start=1):
                    search_in_table(nested_table, f"{cell_path} -> NestedTable {t_index}")

    # --- Body paragraphs
    for i, para in enumerate(doc.paragraphs, start=1):
        highlight_in_paragraph(para, search_text, matches, f"Paragraph {i}")

    # --- Tables
    for t_index, table in enumerate(doc.tables, start=1):
        search_in_table(table, f"Table {t_index}")

    # --- Headers & footers
    for s_index, section in enumerate(doc.sections, start=1):
        for i, para in enumerate(section.header.paragraphs, start=1):
            highlight_in_paragraph(para, search_text, matches,
                                   f"Header Section {s_index}, Paragraph {i}")
        for i, para in enumerate(section.footer.paragraphs, start=1):
            highlight_in_paragraph(para, search_text, matches,
                                   f"Footer Section {s_index}, Paragraph {i}")

    doc.save(output_path)
    return matches


# Example usage
if __name__ == "__main__":
    file_path = "Sampledocx.docx"
    query = "new_tobefound"
    output_file = "highlighted.docx"

    results = search_and_highlight(file_path, query, output_file)

    if results:
        print("Matches found and highlighted:")
        for r in results:
            print(r)
    else:
        print("No match found.")
